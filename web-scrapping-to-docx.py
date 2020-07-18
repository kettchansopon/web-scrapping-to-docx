import requests
from bs4 import BeautifulSoup, SoupStrainer
from PIL import Image
from urllib.request import urlopen
from docx import Document
from docx.shared import Inches
import io

doc = Document()
addHeading = doc.add_heading('Dog Breed Gallery',0).add_run().bold = True

r = requests.get('https://www.pedigree.com.au/dog-breed-information/dog-breed-gallery')
s = BeautifulSoup(r.text, 'lxml')
d = s.find('div', {'id':'dogAZ'})
a_tags = d.find_all('a')

soupStnr = BeautifulSoup(r.text, 'lxml', parse_only=SoupStrainer('img'))
img_tags = soupStnr.find_all('img')
url_prefix = 'https://www.pedigree.com.au'

t = doc.add_table(rows=1, cols=3)  
t.style = "Table Grid"   
fields = t.rows[0].cells
fields[0].text = 'Dog Name'
fields[1].text = 'Link'
fields[2].text = 'Image'

for e in a_tags:
    breed = e.find('img')['alt'] # breed's name
    link = url_prefix + e['href'] # breed's link information
    src = url_prefix + e.find('img')['src'] # image's url
    image_from_url = urlopen(src)
    io_url = io.BytesIO()
    io_url.write(image_from_url.read())
    io_url.seek(0)

    fields = t.add_row().cells
    fields[0].text = breed
    fields[1].text = link
    image_field = fields[2].add_paragraph('').add_run().add_picture(io_url, width = Inches(1.5))

doc.save('document.docx')   
